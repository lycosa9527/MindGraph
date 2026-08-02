"""Build printable learning-sheet DOCX: editable header text + diagram image."""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any, Literal

from docx import Document as create_document
from docx.document import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from docx.image.image import Image as DocxImage
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.styles.style import ParagraphStyle

WorksheetLayout = Literal["landscape", "portrait"]

A4_PORTRAIT_MM = (210.0, 297.0)
A4_LANDSCAPE_MM = (297.0, 210.0)
MARGIN_MM = 15.0
# Header block (topic / meta / instructions) keeps ~38% of content height.
DIAGRAM_HEIGHT_FRACTION = 0.62
BLANK_LINE = "____________"
LATIN_FONT = "Times New Roman"
CJK_FONT = "宋体"
_DOCX_IMAGE_ERRORS = (
    UnrecognizedImageError,
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
)


@dataclass(frozen=True)
class WorksheetDocxLabels:
    """Localized field labels for the DOCX header."""

    name: str
    class_name: str
    date: str
    instruction_prefix: str
    default_instruction: str


@dataclass(frozen=True)
class WorksheetDocxSpec:
    """Worksheet content + placement for DOCX export."""

    title: str
    layout: WorksheetLayout
    show_topic: bool
    show_name: bool
    show_class: bool
    show_date: bool
    show_instruction: bool
    topic_text: str
    instruction_text: str
    diagram_offset_x: float
    diagram_offset_y: float
    diagram_scale: float
    labels: WorksheetDocxLabels


def _clamp_offset(value: float) -> float:
    if value < -1.0:
        return -1.0
    if value > 1.0:
        return 1.0
    return value


def _clamp_scale(value: float) -> float:
    if value < 0.25:
        return 0.25
    if value > 1.0:
        return 1.0
    return value


def _set_r_fonts(r_fonts: Any) -> None:
    """Latin → Times New Roman; Chinese → 宋体 (Word picks by script)."""
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)


def _configure_default_fonts(document: Document) -> None:
    style = document.styles["Normal"]
    if not isinstance(style, ParagraphStyle):
        return
    style.font.name = LATIN_FONT
    style.font.size = Pt(11)
    r_pr = style.element.get_or_add_rPr()
    _set_r_fonts(r_pr.get_or_add_rFonts())


def _stamp_run_fonts(document: Document) -> None:
    """Apply dual fonts to every text run without touching Run protected attrs."""
    for r_elem in document.element.xpath(".//w:r"):
        r_pr = r_elem.get_or_add_rPr()
        _set_r_fonts(r_pr.get_or_add_rFonts())


def _configure_page(document: Document, layout: WorksheetLayout) -> tuple[float, float]:
    section = document.sections[0]
    if layout == "landscape":
        page_w, page_h = A4_LANDSCAPE_MM
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        page_w, page_h = A4_PORTRAIT_MM
        section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(page_w)
    section.page_height = Mm(page_h)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    return page_w - (MARGIN_MM * 2), page_h - (MARGIN_MM * 2)


def _add_topic(document: Document, topic: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(topic)
    run.bold = True
    run.font.size = Pt(16)
    paragraph.paragraph_format.space_after = Pt(10)


def _add_meta_row(document: Document, parts: list[str]) -> None:
    if not parts:
        return
    paragraph = document.add_paragraph("    ".join(parts))
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.size = Pt(11)


def _add_instruction(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(12)
    for run in paragraph.runs:
        run.font.size = Pt(11)


def _fit_diagram_size_mm(
    diagram_bytes: bytes,
    max_width_mm: float,
    max_height_mm: float,
) -> tuple[float, float]:
    """Fit PNG/JPEG into the diagram box while preserving aspect ratio."""
    try:
        image = DocxImage.from_blob(diagram_bytes)
    except _DOCX_IMAGE_ERRORS as exc:
        raise ValueError("diagram image is invalid or unsupported") from exc

    px_width = max(1, int(image.px_width))
    px_height = max(1, int(image.px_height))
    aspect = px_width / px_height
    width_mm = max_width_mm
    height_mm = width_mm / aspect
    if height_mm > max_height_mm:
        height_mm = max_height_mm
        width_mm = height_mm * aspect
    return width_mm, height_mm


def _add_diagram(
    document: Document,
    diagram_bytes: bytes,
    content_width_mm: float,
    content_height_mm: float,
    offset_x: float,
    offset_y: float,
    scale: float,
) -> None:
    ox = _clamp_offset(offset_x)
    oy = _clamp_offset(offset_y)
    sc = _clamp_scale(scale)

    diagram_box_h_mm = max(40.0, content_height_mm * DIAGRAM_HEIGHT_FRACTION)
    max_img_h_mm = diagram_box_h_mm * sc
    max_img_w_mm = content_width_mm * sc
    img_w_mm, img_h_mm = _fit_diagram_size_mm(diagram_bytes, max_img_w_mm, max_img_h_mm)
    free_w = max(0.0, content_width_mm - img_w_mm)
    free_h = max(0.0, diagram_box_h_mm - img_h_mm)
    left_indent_mm = (free_w / 2.0) + (ox * free_w / 2.0)
    space_before_mm = max(0.0, (free_h / 2.0) + (oy * free_h / 2.0))

    if space_before_mm > 0.5:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = Mm(0)
        spacer.paragraph_format.space_after = Mm(space_before_mm)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Mm(left_indent_mm)
    run = paragraph.add_run()
    try:
        run.add_picture(io.BytesIO(diagram_bytes), width=Mm(img_w_mm))
    except _DOCX_IMAGE_ERRORS as exc:
        raise ValueError("diagram image is invalid or unsupported") from exc


def build_worksheet_docx(spec: WorksheetDocxSpec, diagram_bytes: bytes) -> bytes:
    """Return a .docx blob with editable text fields and an embedded diagram PNG/JPEG."""
    if not diagram_bytes:
        raise ValueError("diagram image is required")

    document = create_document()
    _configure_default_fonts(document)
    content_w, content_h = _configure_page(document, spec.layout)
    labels = spec.labels

    topic = spec.topic_text.strip() or spec.title.strip()
    if spec.show_topic and topic:
        _add_topic(document, topic)

    meta_parts: list[str] = []
    if spec.show_name:
        meta_parts.append(f"{labels.name}{BLANK_LINE}")
    if spec.show_class:
        meta_parts.append(f"{labels.class_name}{BLANK_LINE}")
    if spec.show_date:
        meta_parts.append(f"{labels.date}{BLANK_LINE}")
    _add_meta_row(document, meta_parts)

    if spec.show_instruction:
        instruction = spec.instruction_text.strip() or labels.default_instruction
        _add_instruction(
            document,
            f"{labels.instruction_prefix}{instruction}",
        )

    _add_diagram(
        document,
        diagram_bytes,
        content_w,
        content_h,
        spec.diagram_offset_x,
        spec.diagram_offset_y,
        spec.diagram_scale,
    )
    _stamp_run_fonts(document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
