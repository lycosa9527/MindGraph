"""Unit tests for printable learning-sheet DOCX builder."""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm
from PIL import Image

from routers.api.worksheet_docx_export import _content_disposition
from services.diagram.worksheet_docx import (
    A4_LANDSCAPE_MM,
    CJK_FONT,
    DIAGRAM_HEIGHT_FRACTION,
    LATIN_FONT,
    MARGIN_MM,
    WorksheetDocxLabels,
    WorksheetDocxSpec,
    build_worksheet_docx,
)


def _png_bytes(width: int = 80, height: int = 60) -> bytes:
    image = Image.new("RGB", (width, height), color=(32, 64, 128))
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _labels() -> WorksheetDocxLabels:
    return WorksheetDocxLabels(
        name="姓名：",
        class_name="班级：",
        date="日期：",
        instruction_prefix="要求：",
        default_instruction="在空白处补充对应内容",
    )


def _spec(
    *,
    title: str = "水循环",
    layout: str = "landscape",
    show_topic: bool = True,
    show_name: bool = True,
    show_class: bool = True,
    show_date: bool = True,
    show_instruction: bool = True,
    topic_text: str = "水循环",
    instruction_text: str = "",
    diagram_offset_x: float = 0.0,
    diagram_offset_y: float = 0.0,
    diagram_scale: float = 1.0,
) -> WorksheetDocxSpec:
    return WorksheetDocxSpec(
        title=title,
        layout="portrait" if layout == "portrait" else "landscape",
        show_topic=show_topic,
        show_name=show_name,
        show_class=show_class,
        show_date=show_date,
        show_instruction=show_instruction,
        topic_text=topic_text,
        instruction_text=instruction_text,
        diagram_offset_x=diagram_offset_x,
        diagram_offset_y=diagram_offset_y,
        diagram_scale=diagram_scale,
        labels=_labels(),
    )


def test_build_worksheet_docx_includes_editable_fields_and_image() -> None:
    """Editable header labels, default instruction, image, and dual fonts are present."""
    spec = _spec(diagram_scale=0.8)
    docx_bytes = build_worksheet_docx(spec, _png_bytes())
    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "水循环" in text
    assert "姓名：" in text
    assert "班级：" in text
    assert "日期：" in text
    assert "要求：在空白处补充对应内容" in text
    assert document.inline_shapes
    assert len(document.inline_shapes) >= 1
    normal = document.styles["Normal"].element.rPr.rFonts
    assert normal.get(qn("w:ascii")) == LATIN_FONT
    assert normal.get(qn("w:eastAsia")) == CJK_FONT
    stamped = document.element.xpath(".//w:r/w:rPr/w:rFonts")
    assert stamped
    assert stamped[0].get(qn("w:ascii")) == LATIN_FONT
    assert stamped[0].get(qn("w:eastAsia")) == CJK_FONT


def test_content_disposition_allows_chinese_filename() -> None:
    """Content-Disposition stays latin-1 safe and keeps UTF-8 filename*."""
    header = _content_disposition("水循环.docx")
    header.encode("latin-1")
    assert 'filename="worksheet.docx"' in header
    assert "filename*=UTF-8''" in header
    assert "%E6%B0%B4" in header


def test_build_worksheet_docx_hides_disabled_fields() -> None:
    """Hidden worksheet fields are omitted from the document body."""
    spec = _spec(
        title="Topic",
        layout="portrait",
        show_topic=False,
        show_name=False,
        show_class=False,
        show_date=False,
        show_instruction=False,
        topic_text="Hidden",
        instruction_text="Nope",
        diagram_offset_x=1.0,
        diagram_offset_y=-1.0,
        diagram_scale=0.5,
    )
    docx_bytes = build_worksheet_docx(spec, _png_bytes())
    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Hidden" not in text
    assert "姓名：" not in text
    assert "要求：" not in text
    assert document.inline_shapes


def test_build_worksheet_docx_fits_tall_diagram_within_page_box() -> None:
    """Portrait-aspect diagrams shrink to the reserved diagram box height."""
    content_h_mm = A4_LANDSCAPE_MM[1] - (MARGIN_MM * 2)
    max_h_mm = max(40.0, content_h_mm * DIAGRAM_HEIGHT_FRACTION)
    docx_bytes = build_worksheet_docx(_spec(), _png_bytes(width=80, height=400))
    document = Document(io.BytesIO(docx_bytes))
    shape = document.inline_shapes[0]
    assert shape.height <= Mm(max_h_mm)
    assert shape.width <= Mm(A4_LANDSCAPE_MM[0] - (MARGIN_MM * 2))


def test_build_worksheet_docx_rejects_invalid_image() -> None:
    """Corrupt diagram bytes raise ValueError for a clean 400 at the API."""
    with pytest.raises(ValueError, match="invalid or unsupported"):
        build_worksheet_docx(_spec(), b"not-a-real-image")
